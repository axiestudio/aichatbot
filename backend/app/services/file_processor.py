import os
import re
import asyncio
from typing import List, Optional, Dict, Any
from pathlib import Path
import logging
from datetime import datetime

# Document processing imports (with fallbacks)
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

import json

try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False

try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False

# Text processing imports
import tiktoken

try:
    from sentence_transformers import SentenceTransformer
    SENTENCE_TRANSFORMERS_AVAILABLE = True
except ImportError:
    SENTENCE_TRANSFORMERS_AVAILABLE = False

# Create basic document models
from pydantic import BaseModel
from typing import List

class DocumentChunk(BaseModel):
    content: str
    chunk_index: int
    start_char: int
    end_char: int

class DocumentMetadata(BaseModel):
    title: Optional[str] = None
    author: Optional[str] = None

class DocumentType:
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"

logger = logging.getLogger(__name__)


class FileProcessor:
    """Service for processing various file types and extracting content"""
    
    def __init__(self):
        self.encoding = tiktoken.get_encoding("cl100k_base")
    
    async def extract_text(self, file_path: str, file_type: DocumentType) -> str:
        """Extract text content from a file based on its type"""
        try:
            if file_type == DocumentType.PDF:
                return await self._extract_pdf_text(file_path)
            elif file_type == DocumentType.DOCX:
                return await self._extract_docx_text(file_path)
            elif file_type == DocumentType.TXT:
                return await self._extract_txt_text(file_path)
            elif file_type == DocumentType.MD:
                return await self._extract_markdown_text(file_path)
            elif file_type == DocumentType.HTML:
                return await self._extract_html_text(file_path)
            elif file_type == DocumentType.CSV:
                return await self._extract_csv_text(file_path)
            elif file_type == DocumentType.XLSX:
                return await self._extract_xlsx_text(file_path)
            elif file_type == DocumentType.JSON:
                return await self._extract_json_text(file_path)
            else:
                # Fallback to text extraction
                return await self._extract_txt_text(file_path)
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise
    
    async def extract_metadata(self, file_path: str, file_type: DocumentType) -> DocumentMetadata:
        """Extract metadata from a file"""
        try:
            file_stat = os.stat(file_path)
            base_metadata = DocumentMetadata(
                file_size=file_stat.st_size,
                created_date=datetime.fromtimestamp(file_stat.st_ctime),
                modified_date=datetime.fromtimestamp(file_stat.st_mtime)
            )
            
            if file_type == DocumentType.PDF:
                return await self._extract_pdf_metadata(file_path, base_metadata)
            elif file_type == DocumentType.DOCX:
                return await self._extract_docx_metadata(file_path, base_metadata)
            else:
                # For other file types, return basic metadata
                content = await self.extract_text(file_path, file_type)
                base_metadata.word_count = len(content.split())
                base_metadata.char_count = len(content)
                return base_metadata
                
        except Exception as e:
            logger.error(f"Error extracting metadata from {file_path}: {str(e)}")
            return DocumentMetadata()
    
    async def chunk_text(
        self, 
        text: str, 
        chunk_size: int = 1000, 
        chunk_overlap: int = 200,
        document_id: str = None
    ) -> List[DocumentChunk]:
        """Split text into chunks for processing"""
        try:
            # Clean and normalize text
            text = self._clean_text(text)
            
            # Split into sentences for better chunking
            sentences = self._split_into_sentences(text)
            
            chunks = []
            current_chunk = ""
            current_tokens = 0
            chunk_index = 0
            start_char = 0
            
            for sentence in sentences:
                sentence_tokens = len(self.encoding.encode(sentence))
                
                # If adding this sentence would exceed chunk size, create a new chunk
                if current_tokens + sentence_tokens > chunk_size and current_chunk:
                    # Create chunk
                    chunk = DocumentChunk(
                        document_id=document_id,
                        content=current_chunk.strip(),
                        chunk_index=chunk_index,
                        start_char=start_char,
                        end_char=start_char + len(current_chunk),
                        metadata={
                            'token_count': current_tokens,
                            'sentence_count': len(current_chunk.split('. '))
                        }
                    )
                    chunks.append(chunk)
                    
                    # Handle overlap
                    if chunk_overlap > 0:
                        overlap_text = self._get_overlap_text(current_chunk, chunk_overlap)
                        current_chunk = overlap_text + " " + sentence
                        current_tokens = len(self.encoding.encode(current_chunk))
                    else:
                        current_chunk = sentence
                        current_tokens = sentence_tokens
                    
                    start_char += len(chunks[-1].content) - len(overlap_text) if chunk_overlap > 0 else len(chunks[-1].content)
                    chunk_index += 1
                else:
                    # Add sentence to current chunk
                    if current_chunk:
                        current_chunk += " " + sentence
                    else:
                        current_chunk = sentence
                    current_tokens += sentence_tokens
            
            # Add final chunk if there's remaining content
            if current_chunk.strip():
                chunk = DocumentChunk(
                    document_id=document_id,
                    content=current_chunk.strip(),
                    chunk_index=chunk_index,
                    start_char=start_char,
                    end_char=start_char + len(current_chunk),
                    metadata={
                        'token_count': current_tokens,
                        'sentence_count': len(current_chunk.split('. '))
                    }
                )
                chunks.append(chunk)
            
            logger.info(f"Created {len(chunks)} chunks from text")
            return chunks
            
        except Exception as e:
            logger.error(f"Error chunking text: {str(e)}")
            raise
    
    # Private methods for specific file types
    
    async def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            logger.error(f"Error reading PDF {file_path}: {str(e)}")
            raise
        return text
    
    async def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = DocxDocument(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            logger.error(f"Error reading DOCX {file_path}: {str(e)}")
            raise
    
    async def _extract_txt_text(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read()
    
    async def _extract_markdown_text(self, file_path: str) -> str:
        """Extract text from Markdown file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                md_content = file.read()
            # Convert markdown to HTML then extract text
            html = markdown.markdown(md_content)
            soup = BeautifulSoup(html, 'html.parser')
            return soup.get_text()
        except Exception as e:
            logger.error(f"Error reading Markdown {file_path}: {str(e)}")
            raise
    
    async def _extract_html_text(self, file_path: str) -> str:
        """Extract text from HTML file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                html_content = file.read()
            soup = BeautifulSoup(html_content, 'html.parser')
            return soup.get_text()
        except Exception as e:
            logger.error(f"Error reading HTML {file_path}: {str(e)}")
            raise
    
    async def _extract_csv_text(self, file_path: str) -> str:
        """Extract text from CSV file"""
        try:
            df = pd.read_csv(file_path)
            # Convert DataFrame to text representation
            text = f"CSV Data with {len(df)} rows and {len(df.columns)} columns:\n\n"
            text += f"Columns: {', '.join(df.columns)}\n\n"
            text += df.to_string(index=False)
            return text
        except Exception as e:
            logger.error(f"Error reading CSV {file_path}: {str(e)}")
            raise
    
    async def _extract_xlsx_text(self, file_path: str) -> str:
        """Extract text from Excel file"""
        try:
            df = pd.read_excel(file_path)
            # Convert DataFrame to text representation
            text = f"Excel Data with {len(df)} rows and {len(df.columns)} columns:\n\n"
            text += f"Columns: {', '.join(df.columns)}\n\n"
            text += df.to_string(index=False)
            return text
        except Exception as e:
            logger.error(f"Error reading Excel {file_path}: {str(e)}")
            raise
    
    async def _extract_json_text(self, file_path: str) -> str:
        """Extract text from JSON file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
            # Convert JSON to readable text
            return json.dumps(data, indent=2, ensure_ascii=False)
        except Exception as e:
            logger.error(f"Error reading JSON {file_path}: {str(e)}")
            raise
    
    async def _extract_pdf_metadata(self, file_path: str, base_metadata: DocumentMetadata) -> DocumentMetadata:
        """Extract metadata from PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Get document info
                if pdf_reader.metadata:
                    base_metadata.title = pdf_reader.metadata.get('/Title')
                    base_metadata.author = pdf_reader.metadata.get('/Author')
                    base_metadata.subject = pdf_reader.metadata.get('/Subject')
                
                base_metadata.page_count = len(pdf_reader.pages)
                base_metadata.mime_type = "application/pdf"
                
        except Exception as e:
            logger.error(f"Error extracting PDF metadata: {str(e)}")
        
        return base_metadata
    
    async def _extract_docx_metadata(self, file_path: str, base_metadata: DocumentMetadata) -> DocumentMetadata:
        """Extract metadata from DOCX file"""
        try:
            doc = DocxDocument(file_path)
            
            if doc.core_properties:
                base_metadata.title = doc.core_properties.title
                base_metadata.author = doc.core_properties.author
                base_metadata.subject = doc.core_properties.subject
                base_metadata.keywords = doc.core_properties.keywords.split(',') if doc.core_properties.keywords else None
                base_metadata.created_date = doc.core_properties.created
                base_metadata.modified_date = doc.core_properties.modified
            
            # Count words and characters
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + " "
            
            base_metadata.word_count = len(text.split())
            base_metadata.char_count = len(text)
            base_metadata.mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            
        except Exception as e:
            logger.error(f"Error extracting DOCX metadata: {str(e)}")
        
        return base_metadata
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove special characters that might cause issues
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\xff]', '', text)
        # Strip leading/trailing whitespace
        return text.strip()
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences"""
        # Simple sentence splitting - can be enhanced with NLTK or spaCy
        sentences = re.split(r'(?<=[.!?])\s+', text)
        return [s.strip() for s in sentences if s.strip()]
    
    def _get_overlap_text(self, text: str, overlap_size: int) -> str:
        """Get overlap text from the end of a chunk"""
        words = text.split()
        if len(words) <= overlap_size:
            return text
        return ' '.join(words[-overlap_size:])
